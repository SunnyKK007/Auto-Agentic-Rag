import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tech_stack_doc():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Comprehensive Project Architecture & Technology Stack', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph('This document provides an in-depth technical analysis of the architecture and technology stack employed in the Agentic RAG (Retrieval-Augmented Generation) project. It details the purpose of each component, the specific role it plays in the system, and the rationale behind its selection. It also covers the design and justification for two major features: Multi-source Ingestion and Agentic Tool Usage.')
    
    # Frontend Architecture
    doc.add_heading('1. Frontend Architecture: The Interactive UI', level=1)
    doc.add_paragraph('The frontend of this application is responsible for managing user interactions, uploading files to the backend, maintaining the conversation history, and rendering the chat interface.')
    
    doc.add_heading('React.js (Component-Based UI)', level=2)
    doc.add_paragraph('What it does: React is the foundational library used for building the user interface. It manages the component lifecycle and the application state (e.g., the current list of messages, the loading status during API calls, and the currently uploaded files).', style='List Bullet')
    doc.add_paragraph('Why we use it: React’s component-based architecture allows for a highly modular and maintainable codebase. It efficiently updates and renders just the necessary components when data changes, which is crucial for a real-time chat interface where new messages are constantly being appended.', style='List Bullet')
    
    doc.add_heading('Vite (Build Tool & Development Server)', level=2)
    doc.add_paragraph('What it does: Vite is the build tool that compiles the React application and serves it to the browser.', style='List Bullet')
    doc.add_paragraph('Why we use it: Compared to traditional bundlers like Webpack, Vite leverages native ES modules in the browser, resulting in near-instantaneous server startup and lightning-fast Hot Module Replacement (HMR).', style='List Bullet')
    
    doc.add_heading('Tailwind CSS (Utility-First Styling)', level=2)
    doc.add_paragraph('What it does: Tailwind provides low-level utility classes that can be applied directly to HTML elements to style them without writing custom CSS rules.', style='List Bullet')
    doc.add_paragraph('Why we use it: It enables rapid UI development and ensures a consistent design system. By defining styles directly within the React components, we avoid context switching between JS and CSS files.', style='List Bullet')
    
    # Backend Architecture
    doc.add_heading('2. Backend Architecture: Server & API Layer', level=1)
    doc.add_paragraph('The backend acts as the bridge between the frontend user interface and the core AI processing logic. It exposes RESTful APIs to handle requests asynchronously.')
    
    doc.add_heading('FastAPI (Asynchronous Web Framework)', level=2)
    doc.add_paragraph('What it does: FastAPI is a modern Python web framework used to define the API endpoints (/api/ingest, /api/ingest/drive, /api/query).', style='List Bullet')
    doc.add_paragraph('Why we use it: FastAPI is built on Starlette and Pydantic, making it extremely fast. It natively supports asynchronous programming, which is essential for handling multiple concurrent LLM API calls and file processing operations.', style='List Bullet')
    
    # RAG Pipeline
    doc.add_heading('3. Core RAG Pipeline: Data Ingestion & Storage', level=1)
    doc.add_paragraph('This layer is responsible for reading raw files, chunking the text, converting text into mathematical representations, and storing them for rapid semantic retrieval.')
    
    doc.add_heading('LangChain (Orchestration Framework)', level=2)
    doc.add_paragraph('What it does: LangChain provides the standardized abstractions for document loaders, text splitters, and vector store interfaces.', style='List Bullet')
    doc.add_paragraph('Why we use it: It reduces boilerplate code and ensures best practices for building RAG systems.', style='List Bullet')
    
    doc.add_heading('Document Loaders & RecursiveCharacterTextSplitter', level=2)
    doc.add_paragraph('What it does: Parses PDFs, TXTs, CSVs, and chunks them into smaller, overlapping segments (e.g., 800 characters with 10% overlap).', style='List Bullet')
    doc.add_paragraph('Why we use it: LLMs have strict context windows. Chunking ensures we retrieve and pass only the most relevant paragraphs to the LLM.', style='List Bullet')
    
    doc.add_heading('Google Generative AI Embeddings (Gemini)', level=2)
    doc.add_paragraph('What it does: Converts text chunks into high-dimensional vectors to mathematically represent semantic meaning.', style='List Bullet')
    
    doc.add_heading('ChromaDB (Local Vector Database)', level=2)
    doc.add_paragraph('What it does: Stores the generated vectors and text chunks for highly efficient nearest-neighbor searches.', style='List Bullet')
    
    # Agentic Reasoning
    doc.add_heading('4. Agentic Reasoning: The LangGraph Brain', level=1)
    doc.add_paragraph('This is the cognitive core of the application, transforming it from a simple RAG script into an autonomous agent capable of reasoning, planning, and self-correction.')
    
    doc.add_heading('LangGraph (Stateful Workflow Engine)', level=2)
    doc.add_paragraph('What it does: Constructs a cyclical state machine with nodes for planning, retrieving, evaluating relevance, web search fallback, generating answers, and checking hallucinations.', style='List Bullet')
    doc.add_paragraph('Why we use it: It allows the AI to reflect on its intermediate outputs, catch mistakes, and try again, ensuring high accuracy.', style='List Bullet')
    
    doc.add_heading('Google Gemini 2.5 Flash', level=2)
    doc.add_paragraph('What it does: The LLM that executes the prompts at each node in the LangGraph workflow.', style='List Bullet')
    
    # Features
    doc.add_heading('5. Implemented Features: Multi-source & Tools', level=1)
    
    doc.add_heading('Feature 1: Multi-source Ingestion via Google Drive API', level=2)
    doc.add_paragraph('What it does: Integrates google-api-python-client to authenticate with Google Cloud via OAuth 2.0. Users can input a Drive folder ID, Google Doc URL, or Sheet URL. The backend automatically fetches, exports, and ingests the documents into ChromaDB using batch embedding for speed.', style='List Bullet')
    doc.add_paragraph('Why we added this: Bridges the gap between local file uploads and dynamic, cloud-hosted knowledge bases common in enterprise scenarios.', style='List Bullet')
    
    doc.add_heading('Feature 2: Agentic Tool Usage (DuckDuckGo Web Search)', level=2)
    doc.add_paragraph('What it does: Adds a DuckDuckGo web search node to the LangGraph workflow. If the agent evaluates that the local vector database does not contain relevant context (after retries), it autonomously routes to the web search tool to find live internet answers.', style='List Bullet')
    doc.add_paragraph('Why we added this: Elevates the system from a closed-book search engine to a highly autonomous AI assistant capable of answering questions about current events or general knowledge outside the uploaded documents.', style='List Bullet')
    
    doc.save('../Project_Architecture_and_Tech_Stack.docx')
    print("Project Architecture docx created.")


def create_interview_guide_doc():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Agentic RAG Interview Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Core Concepts', level=1)
    
    doc.add_heading('What is RAG (Retrieval-Augmented Generation)?', level=2)
    doc.add_paragraph('RAG is a technique that enhances Large Language Models (LLMs) by giving them access to external data. Instead of relying solely on the data the model was trained on, RAG retrieves relevant information from a custom database (like a vector store) and provides it to the LLM as context to answer a user query. This prevents hallucinations and allows the model to answer questions about private or up-to-date information.')
    
    doc.add_heading('How does "Agentic RAG" differ from Standard RAG?', level=2)
    doc.add_paragraph('Standard RAG is a linear pipeline: Retrieve docs -> Pass to LLM -> Generate answer. If the retrieved docs are irrelevant, the LLM generates a poor answer.')
    doc.add_paragraph('Agentic RAG is a cyclical, stateful system (often using tools like LangGraph). It gives the LLM agency to reason about the process. It can: 1) Formulate its own search queries, 2) Grade the relevance of retrieved documents, 3) Retry searches if documents are irrelevant, 4) Use external tools (like Web Search) if internal knowledge fails, and 5) Verify its own answers for hallucinations before presenting them to the user.')
    
    doc.add_heading('System Architecture', level=1)
    
    doc.add_heading('Why FastAPI?', level=2)
    doc.add_paragraph('FastAPI is built on Python\'s async capabilities (ASGI). In an AI application, you spend a lot of time waiting for external services (LLM APIs, database queries, web scraping). FastAPI handles these concurrent network requests extremely efficiently without blocking the main thread.')
    
    doc.add_heading('What is LangGraph and why use it?', level=2)
    doc.add_paragraph('LangGraph is an extension of LangChain specifically designed for creating cyclical graphs (state machines). We use it to build the cognitive loop of the agent. By defining nodes (Plan, Retrieve, Grade, Web Search, Generate, Check), we can control the exact flow of reasoning and allow the agent to self-correct.')
    
    doc.add_heading('Why use ChromaDB?', level=2)
    doc.add_paragraph('ChromaDB is a local, lightweight vector database. Unlike traditional SQL databases that match exact keywords, ChromaDB stores text as high-dimensional vectors and performs semantic nearest-neighbor searches. It\'s ideal for this project because it persists to disk locally without needing a separate cloud deployment.')
    
    doc.add_heading('Key Features & Problem Solving', level=1)
    
    doc.add_heading('How did you handle Google Drive integration?', level=2)
    doc.add_paragraph('We used the Google Drive API v3 with OAuth 2.0. A major challenge was that Google Workspace files (Docs, Sheets, Slides) cannot be downloaded as raw binary files. We solved this by using the Drive API\'s export endpoint, specifying MIME types (e.g., text/plain for Docs, text/csv for Sheets) to convert them on the fly. We also implemented a URL parser to automatically extract file/folder IDs from user-pasted links.')
    
    doc.add_heading('How did you optimize ingestion speed?', level=2)
    doc.add_paragraph('Initially, ingesting large folders was very slow because the system was making a separate embedding API call for every single text chunk. We solved this by implementing "batch embedding." We group chunks into batches of 10 and send them to the vector store in a single operation, drastically reducing API latency and round-trips.')
    
    doc.add_heading('How do you prevent cross-document data leakage?', level=2)
    doc.add_paragraph('We implemented a session isolation feature. When a user uploads a new document, the frontend passes a "clear_previous" flag. The backend safely invokes the ChromaDB delete_collection() method to wipe the vector store, ensuring the LLM only answers based on the newly uploaded context.')
    
    doc.add_heading('How do you handle missing context?', level=2)
    doc.add_paragraph('We implemented a strict fallback policy. The agent grades retrieved documents. If they are irrelevant, it falls back to DuckDuckGo web search. If web search also fails, or if the final answer check detects a hallucination, the agent strictly outputs: "I cannot find this information in the provided documentation."')
    
    doc.save('../Agentic_RAG_Interview_Guide.docx')
    print("Interview Guide docx created.")

if __name__ == "__main__":
    create_tech_stack_doc()
    create_interview_guide_doc()
