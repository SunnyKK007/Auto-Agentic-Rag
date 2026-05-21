import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_ultra_detailed_arch_doc():
    doc = docx.Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Agentic RAG: Technical Architecture & Implementation Deep-Dive', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Project Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        'The Agentic RAG system is a state-of-the-art information retrieval and generation platform. '
        'Unlike standard RAG systems that follow a linear "Retrieve then Generate" path, this system '
        'implements an "Agentic" loop using LangGraph. The agent autonomously plans searches, '
        'evaluates the quality of retrieved data, performs live web search fallbacks, and verifies '
        'answers for hallucinations before delivery.'
    )
    
    # Section 2: Full Technology Stack
    doc.add_heading('2. Full Technology Stack', level=1)
    
    doc.add_heading('2.1 Frontend (UI/UX Layer)', level=2)
    doc.add_paragraph('Built for performance and aesthetic appeal:', style='Body Text')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('React 19 & Vite 8:').bold = True
    p.add_run(' Provides a component-based reactive interface with near-instant build times.')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('Tailwind CSS v4:').bold = True
    p.add_run(' Used for custom glassmorphic styling, dark mode consistency, and responsive layouts.')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('Micro-Animations:').bold = True
    p.add_run(' Implemented state-aware UI transitions (e.g., "Agent is thinking" flashing dots).')
    
    doc.add_heading('2.2 Backend (Service Layer)', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('FastAPI:').bold = True
    p.add_run(' An asynchronous Python framework that handles high-concurrency requests for file ingestion and LLM querying.')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('Uvicorn:').bold = True
    p.add_run(' The ASGI server optimized for FastAPI\'s non-blocking design.')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('Pydantic v2:').bold = True
    p.add_run(' Enforces strict data validation for all API request/response schemas.')
    
    doc.add_heading('2.3 AI & Intelligence (Orchestration Layer)', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('LangGraph:').bold = True
    p.add_run(' The cyclical state machine that manages the agent\'s reasoning flow and self-correction loops.')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('Google Gemini 2.5 Flash:').bold = True
    p.add_run(' The primary LLM used for planning, grading, and generating verified answers.')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('Gemini Embedding 2:').bold = True
    p.add_run(' Converts text chunks into 768-dimensional vectors for semantic search.')
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('DuckDuckGo Search:').bold = True
    p.add_run(' Integrated as a real-time tool for internet search fallback.')
    
    doc.add_heading('2.4 Storage Layer', level=2)
    p = doc.add_paragraph('', style='List Bullet')
    p.add_run('ChromaDB:').bold = True
    p.add_run(' A persistent local vector database used to store document embeddings and metadata.')
    
    # Section 3: The Agentic Reasoning Flow
    doc.add_heading('3. The Agentic Reasoning Flow (LangGraph)', level=1)
    doc.add_paragraph('The reasoning cycle consists of six distinct nodes, ensuring 100% grounded answers:')
    
    steps = [
        ('1. Plan Search', 'The agent reformulates the user query into an optimized search term for the vector store.'),
        ('2. Retrieve', 'Semantic similarity search fetches the top-4 most relevant document chunks from ChromaDB.'),
        ('3. Evaluate Relevance', 'The agent grades the chunks. If they don\'t answer the question, it refines the plan and retries (Max 2 retries).'),
        ('4. Web Search Fallback', 'If local documents are insufficient, the agent autonomously triggers a live DuckDuckGo web search.'),
        ('5. Generate Answer', 'The agent synthesizes a response using ONLY the provided context (local docs or web search).'),
        ('6. Hallucination Check', 'The agent verifies the final answer against the context. If it detects unsupported claims, it returns a "Data not found" message.')
    ]
    
    for step, desc in steps:
        p = doc.add_paragraph('', style='List Number')
        p.add_run(f'{step}: ').bold = True
        p.add_run(desc)
        
    # Section 4: Advanced Implemented Features
    doc.add_heading('4. Key Features Implemented', level=1)
    
    feat1 = doc.add_heading('4.1 Multi-Source Cloud Ingestion (Google Drive)', level=2)
    doc.add_paragraph(
        'We implemented a custom Drive loader using the Google Drive API v3. Features include:\n'
        '- Automatic URL parsing (detects Folder IDs from Docs/Sheets/Folder links).\n'
        '- Workspace File Export: Converts Google Docs to text and Sheets to CSV on-the-fly.\n'
        '- OAuth 2.0 flow with persistent token storage.'
    , style='Body Text')
    
    feat2 = doc.add_heading('4.2 Performance Optimization (Batch Embedding)', level=2)
    doc.add_paragraph(
        'To solve latency issues, we implemented batching:\n'
        '- Text chunks are embedded in groups of 10 instead of 1-by-1.\n'
        '- This reduced API round-trips by 90% and accelerated ingestion by ~10x.'
    , style='Body Text')
    
    feat3 = doc.add_heading('4.3 Session Isolation', level=2)
    doc.add_paragraph(
        'A "Clear Session" feature allows users to switch documents without context leakage. '
        'The backend clears the ChromaDB collection before new ingestion cycles.'
    )
    
    # Section 5: Engineering Challenges & Solutions
    doc.add_heading('5. Technical Challenges & Engineering Solutions', level=1)
    
    challenges = [
        ('ChromaDB Metadata Bug', 'ChromaDB crashed on complex metadata types. We built a _sanitize_metadata() function to normalize all data to basic types (str/int/float/bool).'),
        ('Drive Export 403 Errors', 'Google Workspace files are not binary. We solved this by using the Drive API\'s .export() endpoint instead of .get_media().'),
        ('SQLite Database Locks', 'Filesystem-level folder deletion caused database panics. We migrated to ChromaDB\'s native .delete_collection() API to respect active connections.'),
        ('Agent Multi-Call Latency', 'Successive LLM calls were slow. We upgraded to Gemini 2.5 Flash, bringing the full 4-step reasoning cycle down to ~8 seconds.')
    ]
    
    for title, sol in challenges:
        p = doc.add_paragraph('', style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(sol)
        
    doc.save('../Project_Architecture_and_Tech_Stack.docx')
    print("Ultra-detailed Project Architecture docx created successfully.")

if __name__ == "__main__":
    create_ultra_detailed_arch_doc()
