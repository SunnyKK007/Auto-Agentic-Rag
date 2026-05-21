import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tech_stack_doc():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Comprehensive Project Architecture & Technology Stack', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph('This document provides an in-depth technical analysis of the architecture and technology stack employed in the Agentic RAG (Retrieval-Augmented Generation) project. It details the purpose of each component, the specific role it plays in the system, and the rationale behind its selection. It also covers the design and justification for the newly implemented advanced features: Google Drive Ingestion, Agentic Web Search, Batch Embedding Optimization, and Session Isolation.')
    
    # Frontend Architecture
    doc.add_heading('1. Frontend Architecture: The Interactive UI', level=1)
    doc.add_paragraph('The frontend of this application is responsible for managing user interactions, uploading files to the backend, maintaining the conversation history, and rendering the chat interface.')
    
    doc.add_heading('React.js (Component-Based UI)', level=2)
    doc.add_paragraph('What it does: React is the foundational library used for building the user interface. It manages the component lifecycle and the application state (e.g., the current list of messages, the loading status during API calls, and the currently uploaded files).', style='List Bullet')
    doc.add_paragraph('Why we use it: React’s component-based architecture allows for a highly modular and maintainable codebase. It efficiently updates and renders just the necessary components when data changes, which is crucial for a real-time chat interface where new messages are constantly being appended.', style='List Bullet')
    
    doc.add_heading('Vite (Build Tool & Development Server)', level=2)
    doc.add_paragraph('What it does: Vite is the build tool that compiles the React application and serves it to the browser.', style='List Bullet')
    doc.add_paragraph('Why we use it: Compared to traditional bundlers like Webpack, Vite leverages native ES modules in the browser, resulting in near-instantaneous server startup and lightning-fast Hot Module Replacement (HMR). This significantly accelerates frontend development.', style='List Bullet')
    
    doc.add_heading('Tailwind CSS v4 (Utility-First Styling)', level=2)
    doc.add_paragraph('What it does: Tailwind provides low-level utility classes that can be applied directly to HTML elements to style them without writing custom CSS rules.', style='List Bullet')
    doc.add_paragraph('Why we use it: It enables rapid UI development and ensures a consistent design system (glassmorphism, dark mode). By defining styles directly within the React components, we avoid context switching between JS and CSS files and eliminate the problem of dead or conflicting CSS code.', style='List Bullet')
    
    # Backend Architecture
    doc.add_heading('2. Backend Architecture: Server & API Layer', level=1)
    doc.add_paragraph('The backend acts as the bridge between the frontend user interface and the core AI processing logic. It exposes RESTful APIs to handle requests asynchronously.')
    
    doc.add_heading('FastAPI (Asynchronous Web Framework)', level=2)
    doc.add_paragraph('What it does: FastAPI is a modern Python web framework used to define the API endpoints: POST /api/ingest, POST /api/ingest/drive, and POST /api/query.', style='List Bullet')
    doc.add_paragraph('Why we use it: FastAPI is built on Starlette and Pydantic, making it one of the fastest Python frameworks available. It natively supports asynchronous programming (async/await), which is essential for handling multiple concurrent LLM API calls and file processing operations without blocking the main server thread.', style='List Bullet')
    
    doc.add_heading('Uvicorn (ASGI Server)', level=2)
    doc.add_paragraph('What it does: Uvicorn is the server that actually runs the FastAPI application, listening for HTTP requests on a specific port.', style='List Bullet')
    doc.add_paragraph('Why we use it: It is a lightning-fast ASGI (Asynchronous Server Gateway Interface) server that perfectly complements FastAPI\'s asynchronous design.', style='List Bullet')
    
    # RAG Pipeline
    doc.add_heading('3. Core RAG Pipeline: Data Ingestion & Storage', level=1)
    doc.add_paragraph('This layer is responsible for reading raw files, chunking the text, converting text into mathematical representations, and storing them for rapid semantic retrieval.')
    
    doc.add_heading('LangChain (Orchestration Framework)', level=2)
    doc.add_paragraph('What it does: LangChain is the overarching framework used to tie all the RAG components together.', style='List Bullet')
    doc.add_paragraph('Why we use it: It provides pre-built, tested components that drastically reduce development time and ensure best practices.', style='List Bullet')
    
    doc.add_heading('Document Loaders & RecursiveCharacterTextSplitter', level=2)
    doc.add_paragraph('What it does: Parses PDFs, TXTs, CSVs from the local filesystem, and divides massive strings of text into smaller, overlapping chunks (e.g., 800 characters per chunk with an 80-character overlap).', style='List Bullet')
    doc.add_paragraph('Why we use it: LLMs have strict context windows. Chunking ensures we retrieve and pass only the most relevant paragraphs to the LLM. The overlap ensures that sentences or concepts are not abruptly cut in half.', style='List Bullet')
    
    doc.add_heading('Google Generative AI Embeddings (Gemini-embedding-2)', level=2)
    doc.add_paragraph('What it does: The embedding model takes a chunk of text and converts it into a high-dimensional vector.', style='List Bullet')
    doc.add_paragraph('Why we use it: By embedding both the query and the document chunks, we can use algorithms (like cosine similarity) to find the chunks that are closest in meaning to the user\'s question.', style='List Bullet')
    
    doc.add_heading('ChromaDB (Local Vector Database)', level=2)
    doc.add_paragraph('What it does: ChromaDB stores the generated vectors along with their corresponding text chunks and metadata.', style='List Bullet')
    doc.add_paragraph('Why we use it: ChromaDB is optimized for highly efficient nearest-neighbor searches in multi-dimensional space. It runs entirely locally, saving data to a local SQLite database, which makes it lightweight, persistent, and perfect for this project\'s scope without needing cloud infrastructure.', style='List Bullet')
    
    # Agentic Reasoning
    doc.add_heading('4. Agentic Reasoning: The LangGraph Brain', level=1)
    doc.add_paragraph('This is the cognitive core of the application, transforming it from a simple RAG script into an autonomous agent capable of reasoning, planning, and self-correction.')
    
    doc.add_heading('LangGraph (Stateful Workflow Engine)', level=2)
    doc.add_paragraph('What it does: LangGraph is used to construct a cyclical state machine. The logic is defined as nodes: Plan, Retrieve, Evaluate Relevance, Web Search (Fallback), Generate Answer, and Check Hallucination.', style='List Bullet')
    doc.add_paragraph('Why we use it: This stateful, cyclical approach is what makes the system "Agentic." It allows the AI to reflect on its own intermediate outputs, catch its own mistakes, and try again, ensuring high accuracy and minimizing hallucinations.', style='List Bullet')
    
    doc.add_heading('Google Gemini 2.5 Flash', level=2)
    doc.add_paragraph('What it does: This is the Large Language Model that executes the prompts at each node in the LangGraph workflow. It grades relevance, checks facts, and synthesizes the final user-facing response.', style='List Bullet')
    doc.add_paragraph('Why we use it: Gemini 2.5 Flash provides an exceptional balance of speed, intelligence, and cost. It is particularly adept at instruction following, which is critical for strict tasks like hallucination grading.', style='List Bullet')
    
    # Advanced Features
    doc.add_heading('5. Advanced Implemented Features', level=1)
    
    doc.add_heading('Feature 1: Google Drive Ingestion (OAuth 2.0 & API v3)', level=2)
    doc.add_paragraph('What it does: Integrates google-api-python-client to allow the backend to authenticate with Google Cloud via OAuth 2.0. Accepts a Google Drive Folder ID, Doc URL, or Sheet URL. Automatically parses the URL, fetches metadata, exports Workspace files to plain text/CSV, downloads PDFs natively, and ingests them into ChromaDB.', style='List Bullet')
    doc.add_paragraph('Why we built this: Real-world enterprise knowledge bases reside in cloud storage, not just local files. This feature bridges that gap, allowing seamless synchronization with dynamic, cloud-hosted folders.', style='List Bullet')
    
    doc.add_heading('Feature 2: Agentic Web Search (DuckDuckGo Fallback)', level=2)
    doc.add_paragraph('What it does: Adds the DuckDuckGoSearchRun tool to the LangGraph workflow. When the LLM determines the local vector database lacks the required context (after retries), the agent autonomously routes the request to Web Search, scrapes live internet results, and formulates an answer based on web context.', style='List Bullet')
    doc.add_paragraph('Why we built this: This elevates the system from a closed-book search engine to a highly autonomous AI assistant capable of answering questions about current events or general knowledge outside the uploaded documents, effectively eliminating dead-ends.', style='List Bullet')
    
    doc.add_heading('Feature 3: Batch Embedding Optimization', level=2)
    doc.add_paragraph('What it does: During ingestion, text chunks are grouped into batches of 10 and sent to the Gemini Embeddings API as a single call. ChromaDB is also updated to insert these batches simultaneously.', style='List Bullet')
    doc.add_paragraph('Why we built this: Processing a large document chunk-by-chunk resulted in extreme latency and API overhead. Batching reduced ingestion time by approximately 10x, making processing of large Google Drive folders feasible.', style='List Bullet')
    
    doc.add_heading('Feature 4: Session Isolation & Context Clearing', level=2)
    doc.add_paragraph('What it does: The frontend provides a "Clear previous session" toggle. When checked, the backend calls vector_store.delete_collection() to safely wipe ChromaDB storage, and the frontend resets chat history before uploading new documents.', style='List Bullet')
    doc.add_paragraph('Why we built this: When users switched between different documents, the vector store retained old context, causing the LLM to mix facts from previously uploaded files into new answers. Session isolation guarantees a completely clean "brain" for every new document analysis.', style='List Bullet')
    
    doc.save('../Project_Architecture_and_Tech_Stack.docx')
    print("Project Architecture docx created successfully.")

if __name__ == "__main__":
    create_tech_stack_doc()
